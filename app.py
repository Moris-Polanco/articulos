import openai
import streamlit as st
import os

# Autenticación de OpenAI (oculta la clave en una variable de entorno)
openai.api_key = os.environ.get("OPENAI_API_KEY")

def generate_article(title, outline):
    model_engine = "text-davinci-003"
    prompt = (f"Escribe un artículo titulado '{title}' basado en el siguiente esquema:\n{outline}")

    completions = openai.Completion.create(
        engine=model_engine,
        prompt=prompt,
        max_tokens=3824,
        n=1,
        stop=None,
        temperature=0.8,
    )

    article = completions.choices[0].text
    return article

st.title("Generador de artículos")

sidebar = st.sidebar
title = sidebar.text_input("Título:")
outline = sidebar.text_area("Esquema o partes:")

if st.button("Generar artículo"):
    article = generate_article(title, outline)
    st.success(article)
if article and st.button("Exportar artículo"):
    from docx import Document
    document = Document()
    document.add_paragraph(article)
    document.save("Article.docx")
    st.success("Artículo exportado en formato docx")

if st.button("Exportar artículo"):
    from docx import Document
    document = Document()
    document.add_paragraph(article)
    document.save("Article.docx")
    st.success("Artículo exportado en formato docx")
